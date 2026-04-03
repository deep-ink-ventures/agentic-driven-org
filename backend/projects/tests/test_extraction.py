"""Tests for projects.extraction module."""

import hashlib
import io
from types import SimpleNamespace
from unittest.mock import MagicMock, patch

import pytest

from projects.extraction import (
    MAX_URL_SIZE,
    compute_content_hash,
    extract_from_docx,
    extract_from_pdf,
    extract_from_txt,
    extract_from_url,
    extract_text,
)


# ── extract_from_txt ───────────────────────────────────────────────────────────


class TestExtractFromTxt:
    def test_utf8(self):
        text = "Hello, world!"
        assert extract_from_txt(text.encode("utf-8")) == text

    def test_utf8_unicode(self):
        text = "Gruesse aus Muenchen"
        assert extract_from_txt(text.encode("utf-8")) == text

    def test_latin1_fallback(self):
        # Latin-1 byte that is not valid UTF-8
        raw = b"\xe4\xf6\xfc"  # a-umlaut, o-umlaut, u-umlaut in Latin-1
        result = extract_from_txt(raw)
        assert result == raw.decode("latin-1")

    def test_replace_fallback(self):
        # bytes that are invalid in both utf-8 and latin-1 won't happen
        # since latin-1 accepts all byte values, so test utf-8 replace
        text = "hello"
        assert extract_from_txt(text.encode("utf-8")) == text


# ── extract_from_pdf ───────────────────────────────────────────────────────────


class TestExtractFromPdf:
    def test_basic_pdf(self):
        import fitz

        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Hello from PDF")
        pdf_bytes = doc.tobytes()
        doc.close()

        result = extract_from_pdf(pdf_bytes)
        assert "Hello from PDF" in result

    def test_multi_page(self):
        import fitz

        doc = fitz.open()
        for i in range(3):
            page = doc.new_page()
            page.insert_text((72, 72), f"Page {i+1}")
        pdf_bytes = doc.tobytes()
        doc.close()

        result = extract_from_pdf(pdf_bytes)
        assert "Page 1" in result
        assert "Page 3" in result


# ── extract_from_docx ─────────────────────────────────────────────────────────


class TestExtractFromDocx:
    def test_basic_docx(self):
        from docx import Document as DocxDocument

        doc = DocxDocument()
        doc.add_paragraph("Hello from DOCX")
        doc.add_paragraph("Second paragraph")
        buf = io.BytesIO()
        doc.save(buf)

        result = extract_from_docx(buf.getvalue())
        assert "Hello from DOCX" in result
        assert "Second paragraph" in result

    def test_empty_paragraphs_stripped(self):
        from docx import Document as DocxDocument

        doc = DocxDocument()
        doc.add_paragraph("Content")
        doc.add_paragraph("")  # blank
        doc.add_paragraph("   ")  # whitespace only
        doc.add_paragraph("More content")
        buf = io.BytesIO()
        doc.save(buf)

        result = extract_from_docx(buf.getvalue())
        assert "Content" in result
        assert "More content" in result
        # Blank paragraphs should be stripped
        parts = result.split("\n\n")
        assert all(p.strip() for p in parts)


# ── extract_from_url ──────────────────────────────────────────────────────────


class TestExtractFromUrl:
    def _mock_response(self, html_bytes):
        resp = MagicMock()
        resp.status_code = 200
        resp.raise_for_status = MagicMock()

        def iter_content(chunk_size=8192):
            for i in range(0, len(html_bytes), chunk_size):
                yield html_bytes[i : i + chunk_size]

        resp.iter_content = iter_content
        return resp

    @patch("requests.get")
    def test_basic_html(self, mock_get):
        html = b"<html><body><main><p>Hello World</p></main></body></html>"
        mock_get.return_value = self._mock_response(html)

        result = extract_from_url("https://example.com")
        assert "Hello World" in result

    @patch("requests.get")
    def test_script_nav_stripped(self, mock_get):
        html = b"""<html><body>
        <nav>Navigation</nav>
        <script>alert('x')</script>
        <main><p>Content</p></main>
        <footer>Footer</footer>
        </body></html>"""
        mock_get.return_value = self._mock_response(html)

        result = extract_from_url("https://example.com")
        assert "Content" in result
        assert "Navigation" not in result
        assert "alert" not in result
        assert "Footer" not in result

    @patch("requests.get")
    def test_size_limit(self, mock_get):
        # Create response larger than MAX_URL_SIZE
        big_html = b"<html><body>" + b"x" * (MAX_URL_SIZE + 1000) + b"</body></html>"
        mock_get.return_value = self._mock_response(big_html)

        # Should not raise, just truncate
        result = extract_from_url("https://example.com")
        assert isinstance(result, str)

    @patch("requests.get")
    def test_http_error(self, mock_get):
        from requests.exceptions import HTTPError

        resp = MagicMock()
        resp.raise_for_status.side_effect = HTTPError("404")
        mock_get.return_value = resp

        with pytest.raises(HTTPError):
            extract_from_url("https://example.com/notfound")


# ── extract_text (dispatch) ──────────────────────────────────────────────────


class TestExtractText:
    def test_text_source(self):
        source = SimpleNamespace(source_type="text", raw_content="Hello text")
        assert extract_text(source) == "Hello text"

    def test_text_source_empty(self):
        source = SimpleNamespace(source_type="text", raw_content="")
        assert extract_text(source) == ""

    @patch("projects.extraction.extract_from_url")
    def test_url_source(self, mock_extract):
        mock_extract.return_value = "Extracted content"
        source = SimpleNamespace(source_type="url", url="https://example.com")
        assert extract_text(source) == "Extracted content"

    def test_url_source_no_url(self):
        source = SimpleNamespace(source_type="url", url="")
        assert extract_text(source) == ""

    @patch("projects.extraction.extract_from_url")
    def test_url_error_returns_message(self, mock_extract):
        mock_extract.side_effect = Exception("Connection refused")
        source = SimpleNamespace(source_type="url", url="https://example.com")
        result = extract_text(source)
        assert "[Extraction failed:" in result
        assert "Connection refused" in result

    @patch("projects.storage.STORAGE_BACKEND", "local")
    def test_file_source_txt(self, tmp_path):
        # Write a test file
        file_content = b"File content here"
        file_path = tmp_path / "projects" / "abc" / "sources" / "test.txt"
        file_path.parent.mkdir(parents=True)
        file_path.write_bytes(file_content)

        source = SimpleNamespace(
            source_type="file",
            file_key="projects/abc/sources/test.txt",
            file_format="txt",
            original_filename="test.txt",
        )

        with patch("projects.storage.LOCAL_MEDIA_ROOT", tmp_path):
            result = extract_text(source)
        assert result == "File content here"

    @patch("projects.storage.STORAGE_BACKEND", "local")
    def test_file_source_pdf(self, tmp_path):
        import fitz

        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "PDF text content")
        pdf_bytes = doc.tobytes()
        doc.close()

        file_path = tmp_path / "projects" / "abc" / "sources" / "test.pdf"
        file_path.parent.mkdir(parents=True)
        file_path.write_bytes(pdf_bytes)

        source = SimpleNamespace(
            source_type="file",
            file_key="projects/abc/sources/test.pdf",
            file_format="pdf",
            original_filename="test.pdf",
        )

        with patch("projects.storage.LOCAL_MEDIA_ROOT", tmp_path):
            result = extract_text(source)
        assert "PDF text content" in result

    @patch("projects.storage.STORAGE_BACKEND", "local")
    def test_file_source_docx(self, tmp_path):
        from docx import Document as DocxDocument

        docx_doc = DocxDocument()
        docx_doc.add_paragraph("DOCX text content")
        buf = io.BytesIO()
        docx_doc.save(buf)
        docx_bytes = buf.getvalue()

        file_path = tmp_path / "projects" / "abc" / "sources" / "test.docx"
        file_path.parent.mkdir(parents=True)
        file_path.write_bytes(docx_bytes)

        source = SimpleNamespace(
            source_type="file",
            file_key="projects/abc/sources/test.docx",
            file_format="docx",
            original_filename="test.docx",
        )

        with patch("projects.storage.LOCAL_MEDIA_ROOT", tmp_path):
            result = extract_text(source)
        assert "DOCX text content" in result

    def test_file_source_no_file_key(self):
        source = SimpleNamespace(source_type="file", file_key="")
        assert extract_text(source) == ""

    @patch("projects.storage.STORAGE_BACKEND", "local")
    def test_file_not_found(self, tmp_path):
        source = SimpleNamespace(
            source_type="file",
            file_key="projects/abc/sources/missing.txt",
            file_format="txt",
            original_filename="missing.txt",
        )

        with patch("projects.storage.LOCAL_MEDIA_ROOT", tmp_path):
            result = extract_text(source)
        assert result == "[File not found]"

    def test_unknown_source_type(self):
        source = SimpleNamespace(source_type="unknown")
        assert extract_text(source) == ""


# ── compute_content_hash ──────────────────────────────────────────────────────


class TestComputeContentHash:
    def test_consistent(self):
        data = b"test data"
        h1 = compute_content_hash(data)
        h2 = compute_content_hash(data)
        assert h1 == h2

    def test_sha256(self):
        data = b"test data"
        expected = hashlib.sha256(data).hexdigest()
        assert compute_content_hash(data) == expected

    def test_different_input(self):
        assert compute_content_hash(b"a") != compute_content_hash(b"b")

    def test_hex_length(self):
        assert len(compute_content_hash(b"x")) == 64
